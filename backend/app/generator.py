# backend/app/ssp_generator.py
import os, tempfile, json
from docx import Document
from app.storage import S3Client

class SSPGenerator:
    def __init__(self, s3_client: S3Client):
        self.s3 = s3_client
        self.bucket = os.getenv("S3_BUCKET")

    async def generate_ssp(self, task_id: str, mapping, user):
        # create a docx with sections per control, include evidence & remediation POA&M
        doc = Document()
        doc.add_heading(f"System Security Plan (CMMC Level 2) - Task {task_id}", level=1)
        doc.add_paragraph(f"Generated for tenant: {user.tenant_id}")
        doc.add_paragraph("Executive summary:")
        doc.add_paragraph("...automatically generated. Please review and sign off.")
        for m in mapping:
            ctrl = m["control"]
            doc.add_heading(f"{ctrl['id']} - {ctrl['title']}", level=2)
            doc.add_paragraph(ctrl["text"])
            doc.add_paragraph(f"Confidence score: {m['score']:.3f}")
            doc.add_paragraph("Evidence:")
            for e in m["evidence"]:
                doc.add_paragraph(f" - chunk {e['chunk_index']} (sim={e['sim']:.3f})")
            # identify gaps (if score below threshold)
            if m["score"] < 0.4:
                doc.add_paragraph("Status: NOT SATISFIED")
                doc.add_paragraph("Recommended POA&M: <insert remediation steps and owner and ETA>")
            else:
                doc.add_paragraph("Status: Likely satisfied (needs review)")

        # save docx
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        s3_key = f"{user.tenant_id}/outputs/{task_id}/SSP.docx"
        self.s3.upload_stream(s3_key, open(tmp.name, "rb").read())
        # also generate PDF using WeasyPrint if you want (omitted for brevity)
        return tmp.name

    async def get_ssp_for_task(self, task_id: str, user):
        # naive: check local tmp (in prod store manifest in DB)
        local = f"/tmp/{task_id}_ssp.docx"
        if os.path.exists(local):
            return local
        return None
